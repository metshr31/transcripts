#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Press Release Tracker
- Fetch releases from configurable RSS feeds
- Filter by WATCHLIST within LOOKBACK_HOURS
- Save results to OUTPUT_DIR in formats listed in OUTPUT_FORMATS (xlsx,csv,docx,pdf,json)
- Email the files using Yahoo SMTP credentials from env vars

Env vars (with sensible defaults):
  OUTPUT_DIR         -> default: "outputs"
  OUTPUT_FORMATS     -> default: "xlsx"   (comma-separated: xlsx,csv,docx,pdf,json)
  MAIL_FROM_NAME     -> default: "Press Release Tracker"
  LOOKBACK_HOURS     -> default: "24"
  SEND_ALWAYS        -> default: "true"   ("true"/"false")
  YAHOO_EMAIL        -> required to send email
  YAHOO_APP_PASSWORD -> required to send email
  TO_EMAIL           -> required to send email
  WATCHLIST          -> default list of TL/LTL/Intermodal & Class I names (comma-separated)
  FEED_URLS          -> optional, comma-separated custom feed URLs

You can override WATCHLIST/FEED_URLS in the workflow env or repo secrets if you like.
"""

from __future__ import annotations
import os
import re
import ssl
import json
import math
import smtplib
import mimetypes
import traceback
from email.message import EmailMessage
from datetime import datetime, timedelta, timezone

import feedparser
import pandas as pd
from bs4 import BeautifulSoup
from dateutil import tz

# ----------------------------- Configuration -----------------------------

DEFAULT_OUTPUT_DIR = "outputs"
DEFAULT_OUTPUT_FORMATS = "xlsx"
DEFAULT_FROM_NAME = "Press Release Tracker"
DEFAULT_LOOKBACK_HOURS = 24
DEFAULT_SEND_ALWAYS = True

DEFAULT_WATCHLIST = [
    # Class I rail
    "Union Pacific", "BNSF", "CSX", "Norfolk Southern", "Canadian National",
    "Canadian Pacific", "CPKC", "Canadian Pacific Kansas City",
    # TL / LTL / 3PL
    "J.B. Hunt", "JB Hunt", "Schneider", "Werner", "Knight-Swift", "Knight Swift",
    "Old Dominion", "ODFL", "Saia", "XPO", "TFI International", "TFI",
    "C.H. Robinson", "CH Robinson", "Landstar", "Hub Group", "RXO", "Uber Freight",
    # Extra common shorthand
    "NS", "UP", "CN"
]

# Reasonable default feeds (you can pass FEED_URLS to override/extend)
DEFAULT_FEEDS = [
    # PR Newswire – master feed
    "https://www.prnewswire.com/rss/news-releases-list.rss",
    # GlobeNewswire – Transportation/Logistics vertical
    "https://www.globenewswire.com/RssFeed/industry/TransportationLogistics/feedTitle/GlobeNewswire%20-%20Transportation%20and%20Logistics",
    # Business Wire – Transportation (RSS param sometimes finicky; ignore errors)
    "https://www.businesswire.com/portal/site/home/news/industry/?vnsId=31350&rss=1&newsLangId=en",
    # Reuters – company/news (broad)
    "http://feeds.reuters.com/reuters/companyNews",
]

# ----------------------------- Helpers: IO -------------------------------

def getenv_str(name: str, default: str | None = None) -> str | None:
    val = os.environ.get(name)
    return val if val is not None and len(val.strip()) > 0 else default

def getenv_bool(name: str, default: bool) -> bool:
    v = os.environ.get(name)
    if v is None:
        return default
    return str(v).strip().lower() in ("1", "true", "yes", "y")

def getenv_list(name: str, default: list[str]) -> list[str]:
    raw = os.environ.get(name)
    if not raw:
        return default[:]
    return [s.strip() for s in raw.split(",") if s.strip()]

def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)

def strip_html(html: str) -> str:
    return BeautifulSoup(html or "", "lxml").get_text(" ", strip=True)

def domain_from_url(url: str) -> str:
    try:
        from urllib.parse import urlparse
        host = urlparse(url).netloc
        return host.lower()
    except Exception:
        return "unknown"

# ----------------------------- Collect ----------------------------------

def _entry_datetime(e) -> datetime | None:
    # feedparser entries sometimes have 'published_parsed' or 'updated_parsed'
    for key in ("published_parsed", "updated_parsed"):
        t = getattr(e, key, None) or e.get(key)
        if t:
            # time.struct_time -> naive UTC
            try:
                return datetime(*t[:6], tzinfo=timezone.utc)
            except Exception:
                pass
    # last resort: try to parse any string fields
    for key in ("published", "updated", "date"):
        s = e.get(key)
        if s:
            try:
                from dateutil import parser as dtparser
                return dtparser.parse(s).astimezone(timezone.utc)
            except Exception:
                continue
    return None

def _match_any(text: str, needles: list[str]) -> list[str]:
    found = []
    low = text.lower()
    for n in needles:
        if n.lower() in low:
            found.append(n)
    return found

def fetch_items(feeds: list[str]) -> list[dict]:
    rows: list[dict] = []
    for url in feeds:
        try:
            feed = feedparser.parse(url)
            for e in feed.entries:
                title = e.get("title", "").strip()
                link = e.get("link", "").strip()
                summary = strip_html(e.get("summary", "") or e.get("description", ""))
                published_dt = _entry_datetime(e)
                src_domain = domain_from_url(link) if link else domain_from_url(feed.get("link", ""))

                rows.append({
                    "source": src_domain or "unknown",
                    "title": title,
                    "url": link,
                    "summary": summary,
                    "published_utc": published_dt.isoformat() if published_dt else "",
                    "published_dt": published_dt,  # keep dt object for sort/filter
                })
        except Exception as ex:
            # Non-fatal; continue other feeds
            print(f"[WARN] Failed feed: {url} — {ex}")
    return rows

def filter_and_tag(rows: list[dict], watchlist: list[str], lookback_hours: int) -> list[dict]:
    now_utc = datetime.now(timezone.utc)
    cutoff = now_utc - timedelta(hours=lookback_hours)
    out = []
    for r in rows:
        dt = r.get("published_dt")
        if dt is None or dt < cutoff:
            continue
        # match watchlist terms in title+summary
        hay = f"{r.get('title','')} {r.get('summary','')}"
        matches = _match_any(hay, watchlist) if watchlist else []
        if watchlist and not matches:
            continue
        r = dict(r)  # shallow copy
        r["companies_matched"] = ", ".join(sorted(set(matches)))
        # Pretty ET time string
        et = tz.gettz("America/New_York")
        if dt:
            r["published_et"] = dt.astimezone(et).strftime("%Y-%m-%d %H:%M %Z")
        else:
            r["published_et"] = ""
        out.append(r)
    # Sort newest first
    out.sort(key=lambda x: x.get("published_dt") or datetime.min.replace(tzinfo=timezone.utc), reverse=True)
    # Drop the dt object before saving
    for r in out:
        r.pop("published_dt", None)
    return out

# ----------------------------- Save (xlsx/csv/docx/pdf/json) -------------

def to_excel(df: pd.DataFrame, path: str) -> None:
    with pd.ExcelWriter(path, engine="openpyxl") as xw:
        df.to_excel(xw, index=False, sheet_name="Results")
        ws = xw.sheets["Results"]
        # autosize columns
        for col_idx, col in enumerate(df.columns, start=1):
            col_values = [str(col)] + [str(v) for v in df[col].tolist()]
            width = min(max(len(v) for v in col_values) + 2, 60)
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width
        ws.freeze_panes = "A2"

def to_docx(df: pd.DataFrame, path: str) -> None:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading("Press Release Tracker", level=1)
    # table
    t = doc.add_table(rows=1, cols=len(df.columns))
    hdr = t.rows[0].cells
    for i, c in enumerate(df.columns):
        p = hdr[i].paragraphs[0]
        p.add_run(str(c)).bold = True
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, c in enumerate(df.columns):
            cells[i].text = str(row[c])
    doc.add_paragraph("")
    doc.add_paragraph(datetime.utcnow().strftime("Generated %Y-%m-%d %H:%M UTC"))
    doc.save(path)

def to_pdf(df: pd.DataFrame, path: str) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    data = [list(df.columns)] + df.astype(str).values.tolist()
    doc = SimpleDocTemplate(path, pagesize=LETTER, title="Press Release Tracker")
    styles = getSampleStyleSheet()
    story = [Paragraph("Press Release Tracker", styles["Title"]), Spacer(1, 12)]
    tbl = Table(data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))
    story.append(Paragraph(datetime.utcnow().strftime("Generated %Y-%m-%d %H:%M UTC"), styles["Normal"]))
    doc.build(story)

def save_results(rows: list[dict], output_dir: str, formats: list[str]) -> list[str]:
    ensure_dir(output_dir)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M")
    df = pd.DataFrame(rows)
    # Column order for readability if present
    preferred_cols = ["source", "companies_matched", "title", "url", "published_et", "published_utc", "summary"]
    cols = [c for c in preferred_cols if c in df.columns] + [c for c in df.columns if c not in preferred_cols]
    if not df.empty:
        df = df[cols]
    paths: list[str] = []
    for fmt in formats:
        fmt = fmt.lower().strip()
        fname = f"press_releases_{ts}.{fmt}"
        fpath = os.path.join(output_dir, fname)
        if fmt == "xlsx":
            to_excel(df, fpath)
        elif fmt == "csv":
            df.to_csv(fpath, index=False)
        elif fmt == "docx":
            to_docx(df, fpath)
        elif fmt == "pdf":
            to_pdf(df, fpath)
        elif fmt == "json":
            df.to_json(fpath, orient="records", indent=2)
        else:
            print(f"[WARN] Unknown format '{fmt}', skipping.")
            continue
        paths.append(fpath)
    return paths

# ----------------------------- Email (Yahoo) -----------------------------

def attach_files(msg: EmailMessage, file_paths: list[str]) -> None:
    for p in file_paths:
        try:
            ctype, _ = mimetypes.guess_type(p)
            maintype, subtype = (ctype or "application/octet-stream").split("/", 1)
            with open(p, "rb") as f:
                msg.add_attachment(f.read(), maintype=maintype, subtype=subtype, filename=os.path.basename(p))
        except Exception as ex:
            print(f"[WARN] Could not attach {p}: {ex}")

def send_email(subject: str, body: str, attachments: list[str], from_name: str) -> None:
    yahoo_email = os.environ.get("YAHOO_EMAIL")
    yahoo_pass = os.environ.get("YAHOO_APP_PASSWORD")
    to_email = os.environ.get("TO_EMAIL")

    if not (yahoo_email and yahoo_pass and to_email):
        print("[INFO] Missing email credentials or TO_EMAIL; skipping email send.")
        return

    msg = EmailMessage()
    msg["From"] = f"{from_name} <{yahoo_email}>"
    msg["To"] = to_email
    msg["Subject"] = subject
    msg.set_content(body)
    attach_files(msg, attachments)

    context = ssl.create_default_context()
    with smtplib.SMTP_SSL("smtp.mail.yahoo.com", 465, context=context) as s:
        s.login(yahoo_email, yahoo_pass)
        s.send_message(msg)
    print(f"[OK] Email sent to {to_email} with {len(attachments)} attachment(s).")

# ----------------------------- Main -------------------------------------

def main() -> int:
    output_dir = getenv_str("OUTPUT_DIR", DEFAULT_OUTPUT_DIR) or DEFAULT_OUTPUT_DIR
    output_formats = getenv_list("OUTPUT_FORMATS", DEFAULT_OUTPUT_FORMATS.split(","))
    from_name = getenv_str("MAIL_FROM_NAME", DEFAULT_FROM_NAME) or DEFAULT_FROM_NAME
    lookback_hours = int(getenv_str("LOOKBACK_HOURS", str(DEFAULT_LOOKBACK_HOURS)))
    send_always = getenv_bool("SEND_ALWAYS", DEFAULT_SEND_ALWAYS)
    watchlist = getenv_list("WATCHLIST", DEFAULT_WATCHLIST)
    feeds = getenv_list("FEED_URLS", DEFAULT_FEEDS)

    print(f"[INFO] lookback_hours={lookback_hours} | formats={output_formats} | output_dir={output_dir}")
    print(f"[INFO] feeds={len(feeds)} | watchlist terms={len(watchlist)}")

    # 1) Fetch
    raw = fetch_items(feeds)
    print(f"[INFO] fetched entries: {len(raw)}")

    # 2) Filter & tag
    rows = filter_and_tag(raw, watchlist, lookback_hours)
    print(f"[INFO] rows after filter: {len(rows)}")

    # 3) Decide whether to send
    if not rows and not send_always:
        print("[INFO] No results and SEND_ALWAYS=false; exiting without output/email.")
        return 0

    # 4) Save files
    files = save_results(rows, output_dir, output_formats)
    print(f"[OK] wrote {len(files)} file(s): {files}")

    # 5) Email
    subject = f"Press Releases – {len(rows)} item(s)"
    if rows:
        # include top few lines in the body
        lines = []
        for r in rows[:10]:
            lines.append(f"- {r.get('published_et','')} | {r.get('source','')} | {r.get('title','')}")
        body = "Here are the latest results.\n\n" + "\n".join(lines)
        if len(rows) > 10:
            body += f"\n… and {len(rows) - 10} more."
    else:
        body = "No items in the selected lookback window."

    send_email(subject=subject, body=body, attachments=files, from_name=from_name)
    return 0

if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except SystemExit:
        raise
    except Exception:
        traceback.print_exc()
        raise